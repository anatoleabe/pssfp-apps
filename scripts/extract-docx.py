#!/usr/bin/env python3
"""
extract-docx.py — Convertit un DOCX en Markdown structuré (cf. spec sprint S5 PR W).

Usage:
    python3 scripts/extract-docx.py <input.docx> [output.md]

Si output.md est omis, le script écrit sur stdout.

Rendu:
    - Paragraphes → texte simple séparé par double retour
    - Heading 1 → `# `, Heading 2 → `## `, etc.
    - Listes → `- ` (bullets) ou `1. ` (numérotées)
    - Bold → `**texte**`, Italic → `*texte*`
    - Lignes vides préservées comme séparateurs

Limites :
    - Pas d'extraction des images (utiliser PR V pour ça).
    - Tableaux convertis en pipe tables Markdown basiques.
"""

import sys
import os
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Erreur : python-docx requis. pip3 install python-docx", file=sys.stderr)
    sys.exit(1)


def style_to_md_prefix(style_name: str) -> str:
    """Map Word heading styles to markdown prefixes."""
    if not style_name:
        return ""
    name = style_name.lower()
    if "heading 1" in name or "titre 1" in name:
        return "# "
    if "heading 2" in name or "titre 2" in name:
        return "## "
    if "heading 3" in name or "titre 3" in name:
        return "### "
    if "heading 4" in name or "titre 4" in name:
        return "#### "
    return ""


def render_runs(runs) -> str:
    """Concat runs with bold/italic markdown."""
    parts = []
    for r in runs:
        text = r.text or ""
        if not text.strip():
            parts.append(text)
            continue
        if r.bold and r.italic:
            parts.append(f"***{text}***")
        elif r.bold:
            parts.append(f"**{text}**")
        elif r.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts)


def is_list_paragraph(p) -> str:
    """Detect list paragraphs (numbered vs bullet)."""
    if not p.style or not p.style.name:
        return ""
    style = p.style.name.lower()
    if "list bullet" in style or "list paragraph" in style:
        return "- "
    if "list number" in style:
        return "1. "
    # Check for numbering format via XML
    pPr = p._p.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr")
    if pPr is not None:
        numPr = pPr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr")
        if numPr is not None:
            return "- "
    return ""


def docx_to_markdown(path: Path) -> str:
    doc = Document(path)
    out_lines: list[str] = []

    # Body title from first paragraph if Heading 1
    for para in doc.paragraphs:
        text = render_runs(para.runs).strip()
        if not text:
            out_lines.append("")
            continue

        h_prefix = style_to_md_prefix(para.style.name if para.style else "")
        list_prefix = is_list_paragraph(para)

        if h_prefix:
            out_lines.append(f"{h_prefix}{text}")
        elif list_prefix:
            out_lines.append(f"{list_prefix}{text}")
        else:
            out_lines.append(text)

    # Append tables as pipe markdown
    for tbl in doc.tables:
        out_lines.append("")
        rows = []
        for r_i, row in enumerate(tbl.rows):
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
            if r_i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        out_lines.extend(rows)
        out_lines.append("")

    # Collapse 3+ blank lines into 2
    cleaned: list[str] = []
    blanks = 0
    for line in out_lines:
        if line.strip() == "":
            blanks += 1
            if blanks <= 2:
                cleaned.append("")
        else:
            blanks = 0
            cleaned.append(line)

    return "\n".join(cleaned).strip() + "\n"


def main() -> int:
    if len(sys.argv) < 2:
        print(__doc__, file=sys.stderr)
        return 2

    src = Path(sys.argv[1])
    if not src.exists():
        print(f"Erreur : fichier introuvable : {src}", file=sys.stderr)
        return 1

    md = docx_to_markdown(src)

    if len(sys.argv) >= 3:
        out_path = Path(sys.argv[2])
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(md, encoding="utf-8")
        print(f"OK : {src} → {out_path} ({len(md)} chars)", file=sys.stderr)
    else:
        sys.stdout.write(md)

    return 0


if __name__ == "__main__":
    sys.exit(main())
